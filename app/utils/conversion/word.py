from docx import Document
import json

# Helper function to convert the input request data into a DOCX
from docx import Document

async def convert_text2docx(lesson_plan_json, output_file_name):
    try:
        # lesson_plan_json is already a Python dict as per your logs
        lesson_plan_data = lesson_plan_json
        
        # Create a new Document object
        doc = Document()

        # Safely extract the list of lessons from the top-level key 'lesson_plan'
        lessons = lesson_plan_data.get("lesson_plan", [])
        
        # Iterate over each lesson (which is a dict)
        for lesson in lessons:
            # Retrieve the "Lesson_Topic" (if any) and add it as a heading
            lesson_topic = lesson.get("Lesson_Topic", "")
            if lesson_topic:
                doc.add_heading(f"Lesson Topic: {lesson_topic}", level=1)
            
            # Now add the remaining details of the lesson
            for key, value in lesson.items():
                # Skip "Lesson_Topic" here because we already added it above
                if key == "Lesson_Topic":
                    continue
                
                # Create a subheading using the key
                doc.add_heading(key.replace("_", " "), level=2)
                
                # Handle the value (could be dict, list, or a simple string)
                if isinstance(value, list):
                    # For lists, create a bullet point for each item
                    for item in value:
                        doc.add_paragraph(f"- {item}")
                elif isinstance(value, dict):
                    # (Optional) If you expect nested dicts, handle them here
                    # For now, just convert dict to string:
                    doc.add_paragraph(str(value))
                else:
                    # Otherwise, just convert it to string
                    doc.add_paragraph(str(value))

        # Define the output file path
        file_path = f"generated_files/{output_file_name}.docx"
        doc.save(file_path)  # Save the DOCX file to the specified path

        return file_path

    except Exception as e:
        # Handle any unforeseen errors
        print(f"Error while generating DOCX: {e}")
        raise
